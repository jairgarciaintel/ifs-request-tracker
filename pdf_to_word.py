#!/usr/bin/env python3
"""Convert the Multi-Party MRUNDA PDF to an editable Word (.docx),
preserving formatting: bold, font sizes, alignment and line structure."""
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = "Multi-Party MRUNDA Acknowledgment (v10-28-2021) FORM.pdf"
OUT = "Multi-Party MRUNDA Acknowledgment (v10-28-2021) FORM.docx"

doc = fitz.open(SRC)
out = Document()
base = out.styles["Normal"]
base.font.name = "Arial"
base.font.size = Pt(11)

def is_bold(font_name):
    f = (font_name or "").lower()
    return "bold" in f or "black" in f

for pnum, page in enumerate(doc):
    pw = page.rect.width
    d = page.get_text("dict")
    # Sort blocks top-to-bottom
    blocks = sorted(d["blocks"], key=lambda b: b.get("bbox", [0, 0])[1])
    for b in blocks:
        if b.get("type", 0) != 0:
            continue  # skip images
        for line in b.get("lines", []):
            spans = line.get("spans", [])
            if not spans:
                continue
            text = "".join(s["text"] for s in spans)
            if not text.strip():
                continue
            p = out.add_paragraph()
            # Alignment: center if the line is roughly centered on the page
            x0 = line["bbox"][0]
            x1 = line["bbox"][2]
            line_center = (x0 + x1) / 2
            if abs(line_center - pw / 2) < 40 and (x1 - x0) < pw * 0.8:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(0)
            for s in spans:
                if not s["text"]:
                    continue
                run = p.add_run(s["text"])
                run.font.name = "Arial"
                run.font.size = Pt(round(s["size"]))
                if is_bold(s["font"]):
                    run.bold = True
                # color
                col = s.get("color", 0)
                r = (col >> 16) & 255
                g = (col >> 8) & 255
                bl = col & 255
                if (r, g, bl) != (0, 0, 0):
                    run.font.color.rgb = RGBColor(r, g, bl)
    if pnum < doc.page_count - 1:
        out.add_page_break()

out.save(OUT)
print("Saved with formatting:", OUT)
